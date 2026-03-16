"""
田中祐貴の分身エージェント — Web版
ブラウザで http://localhost:5050（または Railway の URL）を開いて使う
"""
import io
import json
import os
import queue
import sys
import threading
from pathlib import Path

import anthropic
from flask import Flask, Response, jsonify, render_template, request, stream_with_context

# Windows 日本語対応
if sys.platform == "win32":
    for _name in ("stdout", "stderr"):
        _s = getattr(sys, _name)
        if hasattr(_s, "buffer"):
            setattr(sys, _name, io.TextIOWrapper(_s.buffer, encoding="utf-8", errors="replace"))

# オプション依存ライブラリ
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
try:
    import openpyxl
except ImportError:
    openpyxl = None
try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None
try:
    import requests as req_lib
except ImportError:
    req_lib = None

# ── パス定義 ────────────────────────────────────────────
SCRIPT_DIR   = Path(__file__).parent
PROJECT_DIR  = SCRIPT_DIR.parent
REPO_ROOT    = PROJECT_DIR.parent
DOCS_DIR     = REPO_ROOT / "docs"
PROMPTS_FILE = PROJECT_DIR / "prompts.json"

# ── ナレッジ読み込み ─────────────────────────────────────
def load_knowledge() -> str:
    parts: list[str] = []
    if DOCS_DIR.exists():
        for f in sorted(DOCS_DIR.glob("*.md")):
            try:
                parts.append(f"=== {f.name} ===\n{f.read_text('utf-8')}")
            except Exception:
                pass
    return "\n\n".join(parts)


def load_presets() -> list[dict]:
    default = [{"id": "default", "name": "標準", "prompt": ""}]
    if not PROMPTS_FILE.exists():
        return default
    try:
        data = json.loads(PROMPTS_FILE.read_text("utf-8"))
        presets = data.get("presets") or default
        return [p for p in presets if p.get("id") and p.get("name") is not None]
    except Exception:
        return default


# ── ファイルテキスト抽出 ──────────────────────────────────
def extract_pdf(data: bytes) -> str:
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(p.extract_text() or "" for p in reader.pages)
    except Exception:
        return ""


def extract_docx(data: bytes) -> str:
    if DocxDocument is None:
        return ""
    try:
        doc = DocxDocument(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(parts)
    except Exception:
        return ""


def extract_xlsx(data: bytes) -> str:
    if openpyxl is None:
        return ""
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        for ws in wb.worksheets:
            parts.append(f"[シート: {ws.title}]")
            for row in ws.iter_rows(values_only=True):
                row_str = "\t".join("" if v is None else str(v) for v in row)
                if row_str.strip():
                    parts.append(row_str)
        return "\n".join(parts)
    except Exception:
        return ""


def fetch_url_text(url: str) -> str:
    if req_lib is None:
        return f"（requests未インストールのため取得不可: {url}）"
    try:
        r = req_lib.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
        r.raise_for_status()
        if BeautifulSoup:
            soup = BeautifulSoup(r.content, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            text = soup.get_text(separator="\n")
        else:
            text = r.text
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        return "\n".join(lines)[:8000]
    except Exception as e:
        return f"（取得エラー: {e}）"


# ── システムプロンプト構築 ────────────────────────────────
BASE_PROMPT = """\
あなたは「田中祐貴」の分身AIエージェントです。
田中祐貴は株式会社スリーグッドの営業担当で、X（旧Twitter）アカウント @threee_sales を持ちます。

## できること
1. 営業リサーチ — web_search/web_fetch で企業・業界・競合を調査
2. 提案書作成 — 既存テンプレートを参照してMarkdown提案書を出力
3. トークスクリプト — 架電用スクリプト・Q&Aを作成
4. リスト整理・分析 — CSVの重複排除・集計・ソート
5. スクレイピング — Webサイトから情報収集して構造化
6. Xポスト文作成 — @threee_sales 用の投稿文を複数案作成
7. ブログ記事 — 営業・DX・飲食業向けの専門記事を作成
8. プレスリリース — 会社・サービスのプレスリリース原稿を作成
9. 定例MTG資料 — スプレッドシート用CSVや集計表を出力

## 行動原則
- タスクが与えられたら確認なしに自律的に実行する
- 成果物はチャット上にそのまま出力する
- 調査が必要なときは web_search / web_fetch ツールを積極的に使う
- 田中祐貴のトーン（プロフェッショナル・誠実・具体的な数字）で作成する

## 田中祐貴のペルソナ
- 会社: 株式会社スリーグッド
- X: @threee_sales
- 強み: 飲食店SaaSでアポ率3.0%・HR領域でアポ率1.3〜1.8%
- 人事・採用直通番号リスト20万件以上保有

## 既存ナレッジ
{knowledge}
"""

TOOLS = [
    {"type": "web_search_20260209", "name": "web_search"},
    {"type": "web_fetch_20260209",  "name": "web_fetch"},
]


def build_system_prompt(preset_prompt: str = "", rag_context: str = "") -> str:
    parts = [BASE_PROMPT.format(knowledge=load_knowledge() or "（なし）")]
    if preset_prompt.strip():
        parts.append(f"## 選択中のモード指示\n{preset_prompt.strip()}")
    if rag_context.strip():
        parts.append(f"## アップロードされた参考資料\n{rag_context.strip()}")
    return "\n\n".join(parts)


# ── Claude 呼び出し（別スレッド・SSEキューに積む） ────────
def _run_claude(
    prompt: str,
    preset_prompt: str,
    rag_context: str,
    msg_queue: "queue.Queue[tuple[str,str]]",
) -> None:
    client        = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
    system_prompt = build_system_prompt(preset_prompt, rag_context)
    messages      = [{"role": "user", "content": prompt}]

    try:
        for _ in range(6):          # pause_turn が来た場合に最大6回継続
            with client.messages.stream(
                model="claude-opus-4-6",
                max_tokens=8192,
                thinking={"type": "adaptive"},
                system=system_prompt,
                tools=TOOLS,
                messages=messages,
            ) as stream:
                for text in stream.text_stream:
                    if text:
                        msg_queue.put(("text", text))
                final = stream.get_final_message()

            if final.stop_reason == "end_turn":
                break
            elif final.stop_reason == "pause_turn":
                # サーバー側ツールのループ上限に達した → 続きをリクエスト
                messages.append({"role": "assistant", "content": final.content})
            else:
                break

    except Exception as e:
        msg_queue.put(("error", str(e)))
    finally:
        msg_queue.put(("done", ""))


# ── Flask ────────────────────────────────────────────────
app = Flask(__name__)
app.config["JSON_AS_ASCII"]        = False
app.config["MAX_CONTENT_LENGTH"]   = 50 * 1024 * 1024


@app.route("/")
def index():
    return render_template("agent.html")


@app.route("/api/presets")
def api_presets():
    presets = [{"id": p["id"], "name": p["name"], "prompt": p.get("prompt", "")}
               for p in load_presets()]
    return jsonify({"presets": presets})


@app.route("/api/upload", methods=["POST"])
def api_upload():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "ファイルなし"}), 400
    name  = file.filename or ""
    data  = file.read()
    lower = name.lower()
    if lower.endswith(".pdf"):
        text = extract_pdf(data)
    elif lower.endswith(".docx"):
        text = extract_docx(data)
    elif lower.endswith((".xlsx", ".xls")):
        text = extract_xlsx(data)
    elif lower.endswith((".txt", ".md", ".csv")):
        text = data.decode("utf-8", errors="replace")
    else:
        return jsonify({"error": f"非対応形式: {name}"}), 400
    return jsonify({"name": name, "text": text[:15000]})


@app.route("/api/fetch-url", methods=["POST"])
def api_fetch_url():
    data = request.get_json() or {}
    url  = (data.get("url") or "").strip()
    if not url:
        return jsonify({"error": "URLなし"}), 400
    return jsonify({"url": url, "text": fetch_url_text(url)})


@app.route("/api/chat", methods=["POST"])
def chat():
    data          = request.get_json() or {}
    message       = (data.get("message") or "").strip()
    preset_prompt = (data.get("preset_prompt") or "").strip()
    rag_context   = (data.get("rag_context") or "").strip()
    if not message:
        return jsonify({"error": "メッセージが空です"}), 400

    msg_queue: "queue.Queue[tuple[str,str]]" = queue.Queue()
    threading.Thread(
        target=_run_claude,
        args=(message, preset_prompt, rag_context, msg_queue),
        daemon=True,
    ).start()

    def generate():
        while True:
            try:
                msg_type, content = msg_queue.get(timeout=120)
            except queue.Empty:
                yield "data: [DONE]\n\n"
                break
            if msg_type == "text":
                yield f"data: {json.dumps({'text': content}, ensure_ascii=False)}\n\n"
            elif msg_type == "error":
                yield f"data: {json.dumps({'error': content}, ensure_ascii=False)}\n\n"
                yield "data: [DONE]\n\n"
                break
            elif msg_type == "done":
                yield "data: [DONE]\n\n"
                break

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", os.environ.get("AGENT_PORT", "5050")))
    print(f"http://localhost:{port} を開いてください")
    app.run(host="0.0.0.0", port=port, debug=False, threaded=True)
